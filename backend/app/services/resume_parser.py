"""
Resume Parser Service - 简历解析服务

支持 PDF 和 Word 格式简历的文本提取。
"""

import io
from pathlib import Path
from typing import Optional, Union

import pdfplumber
from docx import Document

from app.core.exceptions import FileProcessingException
from app.utils.logger import get_logger

logger = get_logger(__name__)


class ResumeParser:
    """
    简历解析器
    
    支持 PDF（使用 pdfplumber）和 Word（使用 python-docx）格式。
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}
    
    def __init__(self):
        pass
    
    def parse(
        self,
        file_path: Optional[Union[str, Path]] = None,
        file_content: Optional[bytes] = None,
        file_type: Optional[str] = None,
    ) -> str:
        """
        解析简历文件，提取文本内容
        
        Args:
            file_path: 文件路径（与 file_content 二选一）
            file_content: 文件二进制内容（与 file_content 二选一）
            file_type: 文件类型（当使用 file_content 时必须提供）
        
        Returns:
            提取的文本内容
        
        Raises:
            FileProcessingException: 文件处理失败
        """
        logger.info(f"开始解析简历文件，路径: {file_path}, 类型: {file_type}")
        try:
            if file_path:
                path = Path(file_path)
                ext = path.suffix.lower()
                logger.debug(f"解析文件: {path.name}, 扩展名: {ext}")
                
                if ext not in self.SUPPORTED_EXTENSIONS:
                    logger.warning(f"不支持的文件类型: {ext}")
                    raise FileProcessingException(
                        f"Unsupported file type: {ext}",
                        filename=path.name,
                        file_type=ext,
                    )
                
                if ext == ".pdf":
                    logger.info("解析 PDF 文件")
                    return self._parse_pdf_file(path)
                else:
                    logger.info("解析 Word 文件")
                    return self._parse_docx_file(path)
            
            elif file_content and file_type:
                logger.debug(f"解析字节内容，类型: {file_type}")
                if file_type.lower() in ("pdf", ".pdf", "application/pdf"):
                    logger.info("解析 PDF 字节内容")
                    return self._parse_pdf_bytes(file_content)
                elif file_type.lower() in ("docx", ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
                    logger.info("解析 Word 字节内容")
                    return self._parse_docx_bytes(file_content)
                else:
                    logger.warning(f"不支持的文件类型: {file_type}")
                    raise FileProcessingException(
                        f"Unsupported file type: {file_type}",
                        file_type=file_type,
                    )
            else:
                logger.warning("缺少必要的参数")
                raise FileProcessingException(
                    "Either file_path or (file_content, file_type) must be provided"
                )
        
        except FileProcessingException:
            raise
        except Exception as e:
            logger.error(f"解析简历失败: {str(e)}")
            raise FileProcessingException(
                f"Failed to parse resume: {str(e)}",
                filename=str(file_path) if file_path else None,
            )
    
    def _parse_pdf_file(self, file_path: Path) -> str:
        """解析 PDF 文件"""
        logger.debug(f"开始解析 PDF 文件: {file_path.name}")
        text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            logger.debug(f"PDF 文件页数: {len(pdf.pages)}")
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    logger.debug(f"提取第 {i} 页文本成功")
        
        result = "\n\n".join(text_parts)
        logger.info(f"PDF 文件解析完成，提取文本长度: {len(result)}")
        return result
    
    def _parse_pdf_bytes(self, content: bytes) -> str:
        """解析 PDF 字节内容"""
        logger.debug("开始解析 PDF 字节内容")
        text_parts = []
        
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            logger.debug(f"PDF 字节内容页数: {len(pdf.pages)}")
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    logger.debug(f"提取第 {i} 页文本成功")
        
        result = "\n\n".join(text_parts)
        logger.info(f"PDF 字节内容解析完成，提取文本长度: {len(result)}")
        return result
    
    def _parse_docx_file(self, file_path: Path) -> str:
        """解析 Word 文件"""
        logger.debug(f"开始解析 Word 文件: {file_path.name}")
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 提取表格内容
        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        result = "\n".join(text_parts)
        logger.info(f"Word 文件解析完成，提取文本长度: {len(result)}, 表格数量: {table_count}")
        return result
    
    def _parse_docx_bytes(self, content: bytes) -> str:
        """解析 Word 字节内容"""
        logger.debug("开始解析 Word 字节内容")
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 提取表格内容
        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        result = "\n".join(text_parts)
        logger.info(f"Word 字节内容解析完成，提取文本长度: {len(result)}, 表格数量: {table_count}")
        return result


# 便捷函数
def parse_resume_file(
    file_path: Optional[Union[str, Path]] = None,
    file_content: Optional[bytes] = None,
    file_type: Optional[str] = None,
) -> str:
    """
    解析简历文件的便捷函数
    
    Args:
        file_path: 文件路径
        file_content: 文件二进制内容
        file_type: 文件类型
    
    Returns:
        提取的文本内容
    """
    logger.info("调用 parse_resume_file 函数")
    parser = ResumeParser()
    return parser.parse(file_path=file_path, file_content=file_content, file_type=file_type)
